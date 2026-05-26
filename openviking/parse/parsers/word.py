# Copyright (c) 2026 Beijing Volcano Engine Technology Co., Ltd.
# SPDX-License-Identifier: AGPL-3.0
"""
Word document (.docx) parser for OpenViking.

Converts Word documents to Markdown then parses using MarkdownParser.
Inspired by microsoft/markitdown approach.
"""

import asyncio
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

from openviking.parse.base import ParseResult
from openviking.parse.parsers.base_parser import BaseParser
from openviking.parse.parsers.image_attachments import (
    image_attachment,
    image_extension_from_name_type_or_data,
    image_media_path,
    markdown_image_reference,
)
from openviking_cli.utils.config.parser_config import ParserConfig
from openviking_cli.utils.logger import get_logger

logger = get_logger(__name__)


class WordParser(BaseParser):
    """
    Word document parser for OpenViking.

    Supports: .docx

    Converts Word documents to Markdown using python-docx,
    then delegates to MarkdownParser for tree structure creation.
    """

    def __init__(self, config: Optional[ParserConfig] = None):
        """Initialize Word parser."""
        from openviking.parse.parsers.markdown import MarkdownParser

        self._md_parser = MarkdownParser(config=config)
        self.config = config or ParserConfig()

    @property
    def supported_extensions(self) -> List[str]:
        return [".docx"]

    async def parse(self, source: Union[str, Path], instruction: str = "", **kwargs) -> ParseResult:
        """Parse Word document from file path."""
        path = Path(source)

        if path.exists():
            import docx

            converted = await asyncio.to_thread(self._convert_to_markdown, path, docx)
            if isinstance(converted, tuple):
                markdown_content, attachments = converted
            else:
                markdown_content, attachments = converted, []
            md_kwargs = dict(kwargs)
            attachments = list(md_kwargs.pop("attachments", []) or []) + attachments
            result = await self._md_parser.parse_content(
                markdown_content,
                source_path=str(path),
                instruction=instruction,
                attachments=attachments,
                **md_kwargs,
            )
        else:
            result = await self._md_parser.parse_content(
                str(source), instruction=instruction, **kwargs
            )
        result.source_format = "docx"
        result.parser_name = "WordParser"
        return result

    async def parse_content(
        self, content: str, source_path: Optional[str] = None, instruction: str = "", **kwargs
    ) -> ParseResult:
        """Parse content - delegates to MarkdownParser."""
        result = await self._md_parser.parse_content(content, source_path, **kwargs)
        result.source_format = "docx"
        result.parser_name = "WordParser"
        return result

    def _convert_to_markdown(self, path: Path, docx) -> tuple[str, List[Dict[str, Any]]]:
        """Convert Word document to Markdown string.

        Iterates the document body in order so that tables appear in their
        original position rather than being appended at the end.
        """
        doc = docx.Document(path)
        markdown_parts = []
        attachments: List[Dict[str, Any]] = []
        image_refs: Dict[str, str] = {}
        image_state = {"count": 0}

        # Map XML table elements to python-docx Table objects for O(1) lookup
        table_by_element = {table._tbl: table for table in doc.tables}

        # Walk the document body in order to preserve table positions
        from docx.oxml.ns import qn

        for child in doc.element.body:
            if child.tag == qn("w:p"):
                # It's a paragraph
                from docx.text.paragraph import Paragraph

                paragraph = Paragraph(child, doc)
                text = self._convert_formatted_text(
                    paragraph,
                    doc,
                    attachments,
                    image_refs,
                    image_state,
                )
                if not text.strip():
                    continue

                style_name = paragraph.style.name if paragraph.style else "Normal"

                if style_name.startswith("Heading") and paragraph.text.strip():
                    level = self._extract_heading_level(style_name)
                    markdown_parts.append(f"{'#' * level} {paragraph.text}")
                else:
                    markdown_parts.append(text)

            elif child.tag == qn("w:tbl"):
                # It's a table
                if child in table_by_element:
                    markdown_parts.append(
                        self._convert_table(
                            table_by_element[child],
                            doc,
                            attachments,
                            image_refs,
                            image_state,
                        )
                    )

        return "\n\n".join(markdown_parts), attachments

    def _extract_heading_level(self, style_name: str) -> int:
        """Extract heading level from style name."""
        try:
            if "Heading" in style_name:
                parts = style_name.split()
                for part in parts:
                    if part.isdigit():
                        return min(int(part), 6)
        except Exception:
            pass
        return 1

    def _convert_formatted_text(
        self,
        paragraph,
        doc=None,
        attachments: Optional[List[Dict[str, Any]]] = None,
        image_refs: Optional[Dict[str, str]] = None,
        image_state: Optional[Dict[str, int]] = None,
    ) -> str:
        """Convert paragraph with formatting to markdown."""
        text_parts = []
        for run in paragraph.runs:
            text = run.text
            if text:
                if run.bold:
                    text = f"**{text}**"
                if run.italic:
                    text = f"*{text}*"
                if run.underline:
                    text = f"<ins>{text}</ins>"
                text_parts.append(text)

            if doc is not None and attachments is not None and image_refs is not None:
                text_parts.extend(
                    self._extract_run_images(run, doc, attachments, image_refs, image_state)
                )
        return "".join(text_parts)

    def _convert_table(
        self,
        table,
        doc=None,
        attachments: Optional[List[Dict[str, Any]]] = None,
        image_refs: Optional[Dict[str, str]] = None,
        image_state: Optional[Dict[str, int]] = None,
    ) -> str:
        """Convert Word table to markdown format."""
        if not table.rows:
            return ""

        rows = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_parts = []
                for paragraph in cell.paragraphs:
                    text = self._convert_formatted_text(
                        paragraph,
                        doc,
                        attachments,
                        image_refs,
                        image_state,
                    )
                    if text.strip():
                        cell_parts.append(text.strip())
                row_data.append("<br>".join(cell_parts))
            rows.append(row_data)

        from openviking.parse.base import format_table_to_markdown

        return format_table_to_markdown(rows, has_header=True)

    def _extract_run_images(
        self,
        run,
        doc,
        attachments: List[Dict[str, Any]],
        image_refs: Dict[str, str],
        image_state: Optional[Dict[str, int]],
    ) -> List[str]:
        image_markdown = []
        for rel_id in self._iter_run_image_relationship_ids(run):
            image_part = doc.part.related_parts.get(rel_id)
            if not image_part:
                continue

            part_name = str(getattr(image_part, "partname", rel_id))
            if part_name in image_refs:
                media_path = image_refs[part_name]
            else:
                image_bytes = getattr(image_part, "blob", b"")
                if not image_bytes:
                    continue
                extension = image_extension_from_name_type_or_data(
                    name=part_name,
                    content_type=getattr(image_part, "content_type", ""),
                    data=image_bytes,
                )
                if image_state is None:
                    image_state = {"count": 0}
                image_state["count"] += 1
                media_path = image_media_path(image_state["count"], extension)
                attachments.append(image_attachment(media_path, image_bytes))
                image_refs[part_name] = media_path

            image_markdown.append(markdown_image_reference(media_path))
        return image_markdown

    @staticmethod
    def _iter_run_image_relationship_ids(run) -> List[str]:
        from docx.oxml.ns import qn

        rel_ids: List[str] = []
        for element in run._element.iter():
            if element.tag != qn("a:blip"):
                continue
            rel_id = element.get(qn("r:embed")) or element.get(qn("r:link"))
            if rel_id:
                rel_ids.append(rel_id)
        return rel_ids
